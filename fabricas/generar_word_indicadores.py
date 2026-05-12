"""
Genera Indicadores_CNTXT_PT2.docx
Un bloque por indicador con:
  - Nombre del indicador
  - Cálculo de la medición (procedimiento + período)
  - Línea base (5 decimales)
  - Mes de la medición
  - Meta del indicador (5 decimales)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   'Indicadores_CNTXT_PT2.docx')

# ── Colores ───────────────────────────────────────────────────────────────────
HEAD_RGB  = RGBColor(0x1a, 0x3a, 0x5c)   # azul oscuro
FIJO_RGB  = RGBColor(0xc0, 0x39, 0x2b)   # rojo
VAR_RGB   = RGBColor(0x24, 0x71, 0xa3)   # azul
SOS_RGB   = RGBColor(0x7d, 0x3c, 0x98)   # morado
WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_RGB  = RGBColor(0xEC, 0xF0, 0xF1)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        tag = OxmlElement(f'w:{side}')
        tag.set(qn('w:val'),   kwargs.get('val', 'single'))
        tag.set(qn('w:sz'),    kwargs.get('sz', '6'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', '1a3a5c'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_run(para, text, bold=False, italic=False, size=11,
            color=None, underline=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def heading_para(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    run.font.color.rgb = HEAD_RGB
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1a3a5c')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)
    return p

def label_value_row(doc, label, value, label_color=HEAD_RGB, value_bold=True, value_size=11):
    """Fila etiqueta: valor en párrafo simple."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, f"{label}: ", bold=True, size=10, color=label_color)
    add_run(p, value, bold=value_bold, size=value_size)
    return p

def metric_table(doc, items):
    """
    Tabla de métricas clave: 3 columnas (Línea Base | Mes medición | Meta).
    items = [(label, value, color_hex), ...]
    """
    tbl = doc.add_table(rows=2, cols=len(items))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    for j, (lbl, val, hex_c) in enumerate(items):
        # Fila 0: etiqueta
        cell_lbl = tbl.rows[0].cells[j]
        set_cell_bg(cell_lbl, hex_c)
        p = cell_lbl.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(lbl)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE_RGB
        cell_lbl.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Fila 1: valor
        cell_val = tbl.rows[1].cells[j]
        set_cell_bg(cell_val, 'F7F9FC')
        p2 = cell_val.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(val)
        run2.bold = True
        run2.font.size = Pt(13)
        run2.font.color.rgb = RGBColor(
            int(hex_c[0:2], 16), int(hex_c[2:4], 16), int(hex_c[4:6], 16))
        cell_val.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Altura filas
    for row in tbl.rows:
        row.height = Cm(0.7)
    return tbl

def detail_table(doc, headers, rows, header_hex='1a3a5c'):
    """Tabla de detalle con encabezado coloreado."""
    ncols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Encabezado
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        set_cell_bg(cell, header_hex)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE_RGB

    # Filas
    for i, row_data in enumerate(rows):
        hex_bg = 'FFFFFF' if i % 2 == 0 else 'EAF0FB'
        for j, cell_txt in enumerate(row_data):
            cell = tbl.rows[i+1].cells[j]
            set_cell_bg(cell, hex_bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j <= 1 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(cell_txt))
            run.font.size = Pt(9)
    return tbl

def page_break(doc):
    doc.add_page_break()

def spacer(doc, size=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(size)

def section_divider(doc, text, color_hex, tipo_text):
    """Banda de título de indicador."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'

    cell_l = tbl.rows[0].cells[0]
    cell_r = tbl.rows[0].cells[1]
    set_cell_bg(cell_l, color_hex)
    set_cell_bg(cell_r, color_hex)

    # fusionar celdas no es trivial; usamos 2 celdas con distinto contenido
    p_l = cell_l.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_l = p_l.add_run(f"  {text}")
    run_l.bold = True
    run_l.font.size = Pt(12)
    run_l.font.color.rgb = WHITE_RGB

    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_r = p_r.add_run(f"[ {tipo_text} ]  ")
    run_r.font.size = Pt(9)
    run_r.font.color.rgb = WHITE_RGB

    tbl.rows[0].height = Cm(0.75)
    # anchos
    tbl.columns[0].width = Cm(13)
    tbl.columns[1].width = Cm(4)
    return tbl


# ════════════════════════════════════════════════════════════════════════════
# DATOS DE INDICADORES
# ════════════════════════════════════════════════════════════════════════════
indicadores = [
    dict(
        cod='IPT-1',
        nombre='Reducción de Tiempo que No Agrega Valor en el proceso de gestión de proyectos (TNVA)',
        tipo='Fijo',
        color_hex='C0392B',
        unidad='Minutos por semana',
        base='350,00',
        meta='245,00',
        variacion='-30,00 %',
        mes='Marzo de 2026',
        calculo="""\
Mediante VSM se clasificaron las actividades del proceso de gestión de proyectos (13 activos) en VA y NVA. \
Se midió la frecuencia y duración de cada actividad NVA en una semana tipo de marzo 2026:

NVA 1 – Transcripción manual de datos a tabla para dummies: 3x/sem × 25 min = 75 min.
NVA 2 – Transferencia de datos tabla para dummies a tabla gruesa: 2x/sem × 20 min = 40 min.
NVA 3 – Alimentación de dashboards desde tabla gruesa: 1x/sem × 30 min = 30 min.
NVA 4 – Búsqueda de información dispersa (Drive, Monday, Sheets, correo): 5x/sem × 20 min = 100 min.
NVA 5 – Consolidación manual del reporte ejecutivo semanal: 1x/sem × 45 min = 45 min.
NVA 6 – Actualización manual de cronogramas en Google Sheets: 2x/sem × 15 min = 30 min.
NVA 7 – Reconciliación de costos reales vs. presupuestados: 1x/sem × 30 min = 30 min.

TNVA = 75 + 40 + 30 + 100 + 45 + 30 + 30 = 350 min/semana (línea base, marzo 2026).
Meta: 350 × 0,70 = 245 min/semana (−30 %).""",
        detalle_headers=['Actividad NVA', 'Descripción', 'Cálculo', 'Total (min/sem)'],
        detalle_rows=[
            ('NVA 1', 'Transcripción manual a tabla para dummies', '3x/sem × 25 min', '75'),
            ('NVA 2', 'Transferencia tabla dummies → tabla gruesa', '2x/sem × 20 min', '40'),
            ('NVA 3', 'Alimentación de dashboards desde tabla gruesa', '1x/sem × 30 min', '30'),
            ('NVA 4', 'Búsqueda de información dispersa', '5x/sem × 20 min', '100'),
            ('NVA 5', 'Consolidación manual del reporte ejecutivo', '1x/sem × 45 min', '45'),
            ('NVA 6', 'Actualización de cronogramas Google Sheets', '2x/sem × 15 min', '30'),
            ('NVA 7', 'Reconciliación costos reales vs. presupuesto', '1x/sem × 30 min', '30'),
            ('TOTAL', '', '', '350'),
        ],
        entregable='VSM_CNTXT_PT2.xlsx | VSM_Proceso_Actual_CNTXT.png | VSM_Proceso_Futuro_CNTXT.png',
    ),
    dict(
        cod='IPT-2',
        nombre='Reducción del Tiempo de Ciclo de Gestión Operativa (TCP)',
        tipo='Fijo',
        color_hex='C0392B',
        unidad='Horas hábiles (jornada de 8 horas)',
        base='48,00',
        meta='16,00',
        variacion='-66,67 %',
        mes='Marzo de 2026',
        calculo="""\
La medición se realiza sobre el ciclo completo de respuesta operativa en horas hábiles (jornada de 8 horas), \
desde la ocurrencia de un evento operativo relevante en un proyecto (desfase en cronograma, desviación \
presupuestal, pausa de cliente, demora de disciplina externa) hasta la toma de decisión correctiva informada \
por parte del equipo directivo. El alcance se mantendrá idéntico en las mediciones intermedia y de salida.

Fórmula aplicada: TCP = Σ horas hábiles transcurridas en cada hito del ciclo, desde la ocurrencia del evento \
hasta la decisión correctiva.

Ciclo línea base — 48 horas hábiles (6 días hábiles promedio):

  • Hito 1 — Ocurrencia del evento operativo: Hora 0.
  • Hito 2 — Registro manual del evento por el Líder de Proyecto en Google Sheets:
    12 horas hábiles de latencia (8 a 16 horas si no coincide con día de actualización).
  • Hito 3 — Consolidación de la información en la tabla gruesa:
    12 horas hábiles adicionales.
  • Hito 4 — Preparación del reporte para reunión del equipo directivo (Coffee semanal):
    16 horas hábiles.
  • Hito 5 — Revisión y decisión del equipo directivo durante la reunión semanal:
    4 horas hábiles.
  • Hito 6 — Escalamiento y decisión:
    4 horas hábiles.

TOTAL TCP línea base: 0 + 12 + 12 + 16 + 4 + 4 = 48 horas hábiles (Marzo 2026).
META TCP: 48 × 0,3333 = 16 horas hábiles (reducción del 66,67 %).

Período de la medición: evento operativo tipo representativo del mes de marzo de 2026. \
La temporalidad se expresa en horas hábiles (jornada de 8 horas) conforme a la ficha técnica del indicador.""",
        detalle_headers=['Hito', 'Descripción', 'Línea base', 'Meta'],
        detalle_rows=[
            ('H1', 'Ocurrencia del evento operativo', '0 h', '0 h'),
            ('H2', 'Detección / registro por Líder de Proyecto', '12 h', '1,5 h — alerta automática'),
            ('H3', 'Consolidación en tabla gruesa', '12 h', '0 h — ELIMINADO'),
            ('H4', 'Preparación del reporte para reunión directiva', '16 h', '0 h — ELIMINADO'),
            ('H5', 'Revisión por Líder antes de escalar', '4 h', '3 h — info disponible RT'),
            ('H6', 'Escalamiento y decisión del equipo directivo', '4 h', '11,5 h — decisión informada'),
            ('TOTAL', '', '48 h hábiles', '16 h hábiles'),
        ],
        entregable='VSM_CNTXT_PT2.xlsx | VSM_TCP_CNTXT.png',
    ),
    dict(
        cod='IPT-3',
        nombre='Ahorro Generado en Reducción de Desperdicios Operativos (AD)',
        tipo='Variable',
        color_hex='2471A3',
        unidad='Pesos colombianos por mes',
        base='541.000,00',
        meta='352.000,00',
        variacion='-35,00 %',
        mes='Marzo de 2026',
        calculo="""\
La medición cuantifica en pesos colombianos el costo mensual de las actividades que no agregan valor (TNVA), \
clasificadas en el VSM por perfil involucrado en el proceso de gestión de proyectos. El costo se calcula \
multiplicando las horas TNVA mensuales de cada perfil por su costo por hora real. Base jornada: 182 horas \
laborales mensuales (42 h/semana, Ley 2101 de 2021). El alcance se mantendrá idéntico en las mediciones \
intermedia y de salida.

Fórmula aplicada: AD = Σ (Horas TNVA mensuales del perfil × Costo por hora del perfil).

Las horas TNVA semanales por perfil provienen de la asignación de actividades NVA del VSM, convertidas a \
mensuales con el factor 4,33 sem/mes (4,33 = 52 semanas ÷ 12 meses). \
Fórmula: h/mes = min/sem × 4,33 ÷ 60. \
Los costos por hora se calculan dividiendo honorarios mensuales entre 182 horas (Ley 2101/2021).

  • Perfil 1 — Líder de Proyecto (Mike):
    Actividades NVA: NVA 1 (75 min/sem) + NVA 2 (100 min/sem) + NVA 3 (45 min/sem) + NVA 4 (30 min/sem) = 250 min/sem.
    Mensual: 250 × 4,33 ÷ 60 = 1.082,50 min/mes ÷ 60 = 18,04 h/mes.
    Costo/hora: $4.500.000 ÷ 182 = $24.725/hora.
    Costo TNVA mensual Líder: 18,04 × $24.725 = $446.039/mes.

  • Perfil 2 — Director Ejecutivo (Simón):
    Actividades NVA: NVA 5 (reconciliación manual de costos: 30 min/sem).
    Mensual: 30 × 4,33 ÷ 60 = 129,90 min/mes ÷ 60 = 2,17 h/mes.
    Costo/hora: $8.000.000 ÷ 182 = $43.956/hora.
    Costo TNVA mensual Director Ejecutivo: 2,17 × $43.956 = $95.385/mes.

  • Perfil 3 — Director Creativo (Juan Pablo):
    Actividades NVA asignadas: ninguna (perfil con participación exclusiva en actividades de valor agregado).
    Costo TNVA mensual Director Creativo: $0/mes.

AD línea base: $446.039 + $95.385 + $0 = $541.424/mes → reportado $541.000/mes (Marzo 2026).
Meta AD: $541.000 × 0,65 = $351.650/mes → reportado $352.000/mes (−35 %).
Ahorro mensual esperado: $189.000/mes.""",
        detalle_headers=['Perfil', 'Actividades NVA (min/sem)', 'Horas TNVA/mes', 'Costo/hora (182 h)', 'Costo TNVA/mes'],
        detalle_rows=[
            ('Líder de Proyecto', 'NVA 1+4+5+6  (250 min/sem)', '18,04 h  (250×4,33÷60)', '$24.725', '$446.039'),
            ('Director Ejecutivo', 'NVA 7  (30 min/sem)', '2,17 h  (30×4,33÷60)', '$43.956', '$95.385'),
            ('Director Creativo', 'Ninguna (solo VA)', '0,00 h', '$32.967', '$0'),
            ('TOTAL', '280 min/sem', '20,21 h', '', '$541.424 → reportado $541.000'),
        ],
        entregable='Reporte_Mediciones_CNTXT_PT2.xlsx (hoja IPT-3 AD)',
    ),
    dict(
        cod='IPT-4',
        nombre='Costo Unitario de Gestión Administrativa por Proyecto (CU)',
        tipo='Variable',
        color_hex='2471A3',
        unidad='Pesos colombianos por proyecto por mes',
        base='138.400,00',
        meta='90.000,00',
        variacion='-34,97 %',
        mes='Marzo de 2026',
        calculo="""\
La medición se realiza sobre el costo total de gestión del proceso productivo (actividades que agregan valor \
más actividades que no agregan valor) por cada proyecto activo por mes, para el portafolio de 13 proyectos \
activos simultáneos de CNTXT. El CT incluye todas las actividades de los tres perfiles del proceso según \
tiempos medidos en el VSM. El alcance se mantendrá idéntico en las mediciones intermedia y de salida.

Fórmula aplicada: CU = CT (Costo Total del proceso productivo) ÷ Número de proyectos activos.
CT = suma del costo mensual de todas las actividades de gestión (VA + NVA) de los tres perfiles. Base: 182 h/mes.

Los tiempos por perfil provienen del VSM (minutos por semana). \
Factor de conversión: 4,33 sem/mes (4,33 = 52 semanas ÷ 12 meses). \
Fórmula: h/mes = min/sem × 4,33 ÷ 60. Costos por hora = honorarios ÷ 182 h (Ley 2101/2021).

  • Líder de Proyecto (Mike): 555 min/sem × 4,33 ÷ 60 = 2.403,15 min/mes ÷ 60 = 40,05 h/mes.
    Costo: 40,05 × $24.725 = $990.236/mes.

  • Director Ejecutivo (Simón): 165 min/sem × 4,33 ÷ 60 = 714,45 min/mes ÷ 60 = 11,91 h/mes.
    Costo: 11,91 × $43.956 = $523.516/mes.

  • Director Creativo (Juan Pablo): 120 min/sem × 4,33 ÷ 60 = 519,60 min/mes ÷ 60 = 8,66 h/mes.
    Costo: 8,66 × $32.967 = $285.494/mes.

CT total (VA + NVA): $990.236 + $523.516 + $285.494 = $1.799.246/mes.
Número de proyectos activos: 13 proyectos.
CU = $1.799.246 ÷ 13 = $138.404/proyecto/mes → reportado $138.400/proyecto/mes (Marzo 2026).
Meta CU: $138.400 × 0,65 = $90.000/proyecto/mes (−34,97 %).
Reducción absoluta esperada: $48.400/proyecto/mes.""",
        detalle_headers=['Perfil / Concepto', 'VSM (min/sem)', 'Horas/mes', 'Costo/hora', 'Costo/mes'],
        detalle_rows=[
            ('Líder de Proyecto', '555 min/sem', '40,05 h  (555×4,33÷60)', '$24.725', '$990.236'),
            ('Director Ejecutivo', '165 min/sem', '11,91 h  (165×4,33÷60)', '$43.956', '$523.516'),
            ('Director Creativo', '120 min/sem', '8,66 h  (120×4,33÷60)', '$32.967', '$285.494'),
            ('CT total (VA + NVA)', '', '', '', '$1.799.246 / mes'),
            ('Número de proyectos activos', '', '', '', '13 proyectos'),
            ('CU = $1.799.246 ÷ 13', '', '', '', '$138.404 → $138.400/proy/mes'),
            ('Meta CU (−34,97 %)', '$138.400 × 0,65', '', '', '$90.000/proy/mes'),
            ('Ahorro esperado', '', '', '', '$48.400/proy/mes'),
        ],
        entregable='Reporte_Mediciones_CNTXT_PT2.xlsx (hoja IPT-4 CU)',
    ),
    dict(
        cod='IPT-5',
        nombre='Optimización de Puntos de Contacto con el Cliente en el Proceso Productivo (OPC)',
        tipo='Variable',
        color_hex='2471A3',
        unidad='Porcentaje (%) de subprocesos del ciclo productivo con contacto directo con el cliente externo',
        base='58,33',
        meta='25,00',
        variacion='-57,14 %',
        mes='Marzo de 2026',
        calculo="""\
La medición cuantifica el porcentaje de subprocesos del ciclo productivo (12 actividades del VSM) en los \
cuales se tiene contacto directo con el cliente externo, desde el registro del dato operativo hasta la \
aprobación del hito. Se consideran únicamente los subprocesos que requieren interacción directa con el \
cliente externo (no transferencias internas de información). El alcance se mantendrá idéntico en las \
mediciones intermedia y de salida.

Fórmula aplicada: OPC = (Subprocesos con contacto directo con cliente externo ÷ Total subprocesos del ciclo) × 100.

Total de subprocesos del ciclo productivo: 12 actividades (según VSM_CNTXT_PT2.xlsx).

Línea base — 7 subprocesos con contacto directo con cliente externo = 7/12 = 58,33 % (Marzo 2026):

  • Punto 1: Convocatoria y coordinación de reunión de avance con el cliente (correo o llamada).
  • Punto 2: Preparación de presentación de avance impresa para el cliente.
  • Punto 3: Reunión presencial de avance con el cliente con entrega de material impreso.
  • Punto 4: Envío de planos y renders al cliente para revisión y aprobación (físico o correo).
  • Punto 5: Recepción y procesamiento de comentarios del cliente (verbal o correo).
  • Punto 6: Ciclo de ajustes por feedback del cliente y re-envío de entregables.
  • Punto 7: Confirmación de aprobación del cliente y escalamiento de hito.

Meta — 3 subprocesos con contacto directo con cliente = 3/12 = 25,00 %:

  • Punto 1: Reunión presencial de avance (conservada — exigencia del cliente).
  • Punto 2: Acceso del cliente al dashboard en tiempo real en Central Contexto 2.0.
  • Punto 3: Aprobación digital del cliente en la plataforma (conserva la confirmación formal).

OPC línea base: 7/12 = 58,33 %. Meta: 3/12 = 25,00 %. Variación: −57,14 %.""",
        detalle_headers=['Punto', 'Descripción (contacto directo con cliente externo)', 'Estado en meta'],
        detalle_rows=[
            ('P 1', 'Convocatoria y coordinación de reunión de avance con cliente', 'ELIMINADO — dashboard digital'),
            ('P 2', 'Preparación de presentación de avance impresa para el cliente', 'ELIMINADO — dashboard digital'),
            ('P 3', 'Reunión presencial de avance con cliente con material impreso', 'Se conserva — exigencia cliente'),
            ('P 4', 'Envío de planos y renders al cliente para aprobación', 'ELIMINADO — portal digital'),
            ('P 5', 'Recepción y procesamiento de comentarios del cliente', 'ELIMINADO — formulario digital'),
            ('P 6', 'Ciclo de ajustes por feedback del cliente y re-envío', 'ELIMINADO — trazabilidad plataforma'),
            ('P 7', 'Confirmación de aprobación y escalamiento de hito con el cliente', 'Se conserva — aprobación digital'),
            ('', 'Línea base: 7 / 12 subprocesos = 58,33 %', ''),
            ('', 'Meta: 3 / 12 subprocesos = 25,00 %   →   Variación: −57,14 %', ''),
        ],
        entregable='Reporte_Mediciones_CNTXT_PT2.xlsx (hoja IPT-5 OPC)',
    ),
    dict(
        cod='IPT-6',
        nombre='Impacto Ambiental: Consumo de Recursos por Servicio Prestado (IA)',
        tipo='Variable — Sostenibilidad Ambiental',
        color_hex='7D3C98',
        unidad='Hojas de papel por mes',
        base='170,00',
        meta='145,00',
        variacion='-14,71 %',
        mes='Marzo de 2026',
        calculo="""\
La medición se realiza sobre el consumo mensual de papel asociado a procesos documentales de gestión \
de proyectos en CNTXT, para el portafolio de 13 proyectos activos. Se realizó un inventario físico \
de los tipos de documentos impresos mensualmente y sus cantidades. \
El alcance se mantendrá idéntico en las mediciones intermedia y de salida.

Fórmula aplicada: IA = Σ (Frecuencia mensual × Número de hojas por tipo de documento).
El resultado se expresa en hojas de papel por mes.

  • Documento 1 — Reportes ejecutivos impresos para reuniones presenciales con equipo directivo:
    40 hojas/mes (4 reuniones × 10 hojas promedio).

  • Documento 2 — Presentaciones de avance de proyectos impresas para reuniones con clientes:
    60 hojas/mes (distribuidas entre 13 proyectos activos).

  • Documento 3 — Planos y renders impresos para aprobación de clientes en reuniones presenciales:
    30 hojas/mes.

  • Documento 4 — Documentación de soporte de gestión de proyectos (actas de seguimiento, fichas de entregables):
    25 hojas/mes.

  • Documento 5 — Copias de respaldo de información crítica de proyectos:
    15 hojas/mes.

TOTAL IA línea base: 40 + 60 + 30 + 25 + 15 = 170 hojas/mes (Marzo 2026).
META IA: 145 hojas/mes (reducción del 14,71 %).

Nota: el Documento 4 presenta una meta de 30 hojas (vs. 25 en línea base) debido a que \
la implementación del módulo de prerrequisitos de Central Contexto 2.0 genera inicialmente \
un incremento controlado en actas digitales que aún requieren validación física; \
la reducción neta total se mantiene en 25 hojas/mes gracias a la eliminación de los reportes ejecutivos impresos.

Período de la medición: mes de marzo de 2026. Inventario físico levantado durante la primera semana \
de marzo de 2026. La temporalidad se expresa mensualmente conforme a la ficha técnica del indicador.""",
        detalle_headers=['Documento', 'Descripción', 'Línea base', 'Meta', 'Reducción / Acción'],
        detalle_rows=[
            ('Doc 1', 'Reportes ejecutivos para equipo directivo', '40 h/mes', '15 h/mes', '−25 h. Dashboard digital'),
            ('Doc 2', 'Presentaciones de avance para clientes', '60 h/mes', '55 h/mes', '−5 h.  PDF digital parcial'),
            ('Doc 3', 'Planos y renders para aprobación cliente', '30 h/mes', '30 h/mes', '0 h. Req. cliente físico'),
            ('Doc 4', 'Soporte gestión (actas, fichas)', '25 h/mes', '30 h/mes', '+5 h. Módulo prerreqs.'),
            ('Doc 5', 'Copias de respaldo información crítica', '15 h/mes', '15 h/mes', '0 h. Política interna'),
            ('TOTAL', '', '170 hojas/mes', '145 hojas/mes', '−25 hojas/mes (−14,71 %)'),
        ],
        entregable='Inventario_Papel_CNTXT_PT2_IA.xlsx',
    ),
]


# ════════════════════════════════════════════════════════════════════════════
# CONSTRUIR DOCUMENTO
# ════════════════════════════════════════════════════════════════════════════
doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Portada / encabezado general ──────────────────────────────────────────
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p_title, 'MEDICIÓN DE INDICADORES — PLAN DE TRABAJO 2\n',
        bold=True, size=16, color=HEAD_RGB)
add_run(p_title, 'CONTEXTO ARQUITECTURA S.A.S. (CNTXT Casa de Diseño)\n',
        bold=True, size=12, color=HEAD_RGB)
add_run(p_title, 'Línea de Servicio: Productividad Operacional  |  NIT: 900.948.892-7\n',
        size=10, color=RGBColor(0x5d, 0x6d, 0x7e))
add_run(p_title, 'Consultor: Miguel Bernardo Rodríguez Torres — INDUNNOVA S.A.S.\n',
        size=10, color=RGBColor(0x5d, 0x6d, 0x7e))
add_run(p_title, 'Período intervención: 09 de abril – 18 de junio de 2026  |  Período línea base: Marzo de 2026',
        size=10, color=RGBColor(0x5d, 0x6d, 0x7e))

doc.add_paragraph()

# ── Portafolio de 13 proyectos activos ───────────────────────────────────
heading_para(doc, 'Portafolio de 13 Proyectos Activos — Alcance de la Medición (Marzo 2026)')

p_nota = doc.add_paragraph()
p_nota.paragraph_format.space_after = Pt(6)
add_run(p_nota,
    'Los indicadores IPT-1, IPT-2, IPT-3, IPT-4 e IPT-6 se midieron sobre el portafolio completo '
    'de 13 proyectos activos simultáneos de CNTXT en marzo de 2026. '
    'El alcance se mantendrá idéntico en las mediciones intermedia y de salida.',
    size=10)

PROYECTOS_LIST = [
    ('1',  'Made',   'POP-UP AAA',                     'Comercial',                   'Diseño'),
    ('2',  'Made',   'WELLNESS RESIDENCES',             'Parcelaciones y Condominios', 'Diseño & Visual'),
    ('3',  'Made',   'LA SENDA',                        'Parcelaciones y Condominios', 'Diseño & Visual'),
    ('4',  'Made',   'NATIVE SAN JOSÉ',                 'Parcelaciones y Condominios', 'Diseño & Visual'),
    ('5',  'Made',   'DON MATIAS',                      'Parcelaciones y Condominios', 'Diseño'),
    ('6',  'Made',   'CASA RODINA',                     'Vivienda Campestre',          'Diseño & Visual'),
    ('7',  'Made',   'CASA ESCULTA',                    'Vivienda Campestre',          'Diseño & Visual'),
    ('8',  'Select', 'CASA DIANA & HERIBERTO',          'Vivienda Campestre',          'Diseño & Visual'),
    ('9',  'Select', 'CASA ADRIANA MOCCIOLA & FELIPE',  'Vivienda Campestre',          'Diseño & Visual'),
    ('10', 'Select', 'CASA CARLOS & YOHANA',            'Vivienda Campestre',          'Diseño & Visual'),
    ('11', 'Select', 'CASA CARMEN & TONY',              'Vivienda Campestre',          'Diseño & Visual'),
    ('12', 'Select', 'CASA MAURICIO MARTINEZ',          'Vivienda Campestre',          'Diseño & Visual'),
    ('13', 'Select', 'CASA FELIPE SALDARRIAGA',         'Vivienda Campestre',          'Diseño & Visual'),
]

detail_table(doc,
    headers=['#', 'Modelo', 'Proyecto', 'Categoría', 'Línea Operativa'],
    rows=PROYECTOS_LIST,
    header_hex='1a3a5c',
)
spacer(doc, 8)
page_break(doc)

# ── Un bloque por indicador ───────────────────────────────────────────────
for idx, ind in enumerate(indicadores):

    if idx > 0:
        page_break(doc)

    # ── 1. Banda de título ────────────────────────────────────────────────
    section_divider(doc, f"{ind['cod']}  —  {ind['nombre']}", ind['color_hex'], ind['tipo'])
    spacer(doc, 6)

    # ── 2. Nombre del indicador ───────────────────────────────────────────
    heading_para(doc, '1. Nombre del Indicador')
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    add_run(p, ind['nombre'], bold=True, size=11, color=HEAD_RGB)
    spacer(doc, 4)

    # ── 3. Cálculo de la medición ─────────────────────────────────────────
    heading_para(doc, '2. Cálculo de la Medición')
    for linea in ind['calculo'].split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Cm(0.5)
        add_run(p, linea, size=10)
    spacer(doc, 4)

    # Tabla de detalle del cálculo
    tbl = detail_table(doc, ind['detalle_headers'], ind['detalle_rows'],
                       header_hex=ind['color_hex'])
    spacer(doc, 6)

    # ── 4. Métricas clave (tabla visual) ──────────────────────────────────
    heading_para(doc, '3. Métricas del Indicador')
    spacer(doc, 2)

    metric_table(doc, [
        ('Línea Base', ind['base'], 'C0392B'),
        ('Mes de la Medición', ind['mes'], '1A3A5C'),
        ('Meta del Indicador', ind['meta'], '27AE60'),
        ('Variación Esperada', ind['variacion'], ind['color_hex']),
        ('Unidad de Medida', ind['unidad'], '5D6D7E'),
    ])
    spacer(doc, 6)

    # ── 5. Entregable de soporte ──────────────────────────────────────────
    heading_para(doc, '4. Entregable de Soporte')
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, ind['entregable'], size=10, italic=True,
            color=RGBColor(0x15, 0x43, 0x60))

doc.save(OUT)
import os
size_kb = os.path.getsize(OUT) / 1024
print(f'✓  {os.path.basename(OUT)}  —  {size_kb:.1f} KB')
